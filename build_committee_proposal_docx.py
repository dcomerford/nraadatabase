"""Under-two-page committee proposal for the MCSI V5, for circulation at the meeting.
Deliberately compact — the deep detail lives in the technical response. Same house
style as build_gerald_response_docx.py.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/dancomerford/Desktop/claude/nraadatabase/GRC_MCSI_V5_Committee_Proposal.docx'

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
HEADER_ROW_FILL = 'EEF3F8'
ROW_HIGHLIGHT = 'FFF9E6'
GREY_FILL = 'F4F4F4'

doc = Document()
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _left_border(cell, hex_color, sz='24'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    borders.append(left)
    tcPr.append(borders)


def _add_runs(paragraph, text):
    text = text.replace('&amp;', '&')
    pattern = re.compile(r'(<b>.*?</b>|<i>.*?</i>)', re.S)
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith('<b>') and part.endswith('</b>'):
            paragraph.add_run(part[3:-4]).bold = True
        elif part.startswith('<i>') and part.endswith('</i>'):
            paragraph.add_run(part[3:-4]).italic = True
        else:
            paragraph.add_run(part)


def add_para(text, *, size=10, bold=False, italic=False, color=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if color is not None:
            r.font.color.rgb = color
    return p


def add_title(text):
    add_para(text, size=17, bold=True, color=NAVY, space_after=2)


def add_subtitle(text):
    add_para(text, size=10.5, color=GREY, space_after=5)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = NAVY


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10)


def add_small(text, italic=True):
    add_para(text, size=8.5, italic=italic, color=LIGHT_GREY, space_after=6)


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def add_callout(text, fill=HEADER_ROW_FILL, border=NAVY):
    border_hex = border if isinstance(border, str) else '1F4E78'
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(17.4)
    _shade_cell(cell, fill)
    _left_border(cell, border_hex)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_table(data, col_widths_cm=None, highlight_rows=None):
    n_rows, n_cols = len(data), len(data[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    table.style = 'Table Grid'
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for r_idx, row_data in enumerate(data):
        is_header = r_idx == 0
        is_highlight = highlight_rows and r_idx in highlight_rows
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if col_widths_cm and c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_runs(p, str(text))
            for r in p.runs:
                r.font.size = Pt(9.5)
                if is_header:
                    r.bold = True
                    r.font.color.rgb = NAVY
            if is_header:
                _shade_cell(cell, HEADER_ROW_FILL)
            elif is_highlight:
                _shade_cell(cell, ROW_HIGHLIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ── Header ───────────────────────────────────────────────────────────────────
add_title('Proposal — Adopt the V5 Mixed Category Score Index (MCSI)')
add_subtitle('Geelong Rifle Club · for the committee meeting · July 2026 · one-page summary '
             '(full technical detail available on request)')
add_hr()

# ── The proposal ───────────────────────────────────────────────────────────────
add_h2('What is being proposed')
add_para(
    'That the club adopt the <b>V5 MCSI factors</b> (table below) for the combined club '
    'championship. V5 replaces the older factors with values calibrated on recent competition '
    'data, and — for the first time — includes a properly calibrated <b>Sporter</b> factor '
    'rather than an estimate.')

add_h2('What the MCSI is — and is not')
add_para(
    'The MCSI converts scores shot under different systems (50-point vs 60-point) and different '
    'equipment onto a <b>common scale</b>, so the disciplines can be compared fairly in the '
    'combined championship. It is a <b>constructed comparison index</b> — the same idea as golf '
    'handicaps or decathlon points. It does <b>not</b> claim to identify the objectively best '
    'shooter, and it does not rank a club member against a national champion. Each discipline’s '
    'own champion is still recognised separately.')

add_h2('The formula')
add_callout('<b>Adjusted MCSI = ( (Score × Conversion) + (Centres × 0.7) ) × Factor</b>',
            fill=GREY_FILL, border='444444')
add_para(
    '<i>Conversion</i> puts every discipline on a common base; <i>Centres × 0.7</i> rewards a '
    'well-centred shoot (about 1½ centres make up one dropped point); <i>Factor</i> balances the '
    'disciplines. It stays simple enough to calculate by hand.', space_after=4)

add_h2('The proposed factors')
add_table([
    ['Discipline', 'Scores used', 'Shooters', 'Top-40% average', 'V5 factor'],
    ['Target Rifle', '7,510', '351', '66.9', '<b>1.412</b>'],
    ['F-Open', '4,082', '222', '67.2', '<b>1.406</b>'],
    ['F-Standard', '3,731', '193', '64.1', '<b>1.475</b>'],
    ['F/TR', '2,437', '124', '65.2', '<b>1.450</b>'],
    ['Sporter (Open + Production)', '2,018', '152', '68.3', '<b>1.383</b>'],
], col_widths_cm=[5.6, 3.0, 2.6, 3.4, 2.8])
add_small('~19,800 King’s/Queen’s range scores across 135 championship days (2024 onward). '
          'Sporter Open and Production share one factor for now — the records don’t yet separate '
          'them cleanly.')

add_h2('How each factor is set (no magic numbers)')
add_para(
    'Each factor is <b>one division</b>: it lifts that discipline’s top-40% average onto a single '
    'shared level (94.5, the level the F-classes already sit at). The discipline whose strong '
    'scores average lowest gets the largest factor; the highest gets the smallest. Anyone can '
    'reproduce the table from the averages above.')

add_h2('What V5 changes')
add_bullet('<b>Recent data</b> — calibrated on 2024-onward results, so it reflects today’s '
           'equipment, not a decade ago.')
add_bullet('<b>Transparent method</b> — factors are a published calculation, not chosen by eye.')
add_bullet('<b>Sporter included properly</b> — calibrated from real data for the first time, '
           'replacing an estimate.')

add_h2('Governance — so results stay comparable')
add_para(
    'Factors are <b>locked before each season</b>, reviewed once a year after the season, tested '
    'against past results before any change, and only changed with committee approval — applied '
    'from the <b>next</b> season. Past results are never rewritten, and every factor set is '
    'versioned.')

add_hr()
add_para(
    '<b>Recommendation:</b> adopt the V5 factors above for the coming season under the governance '
    'rule stated. Fuller technical detail — sample sizes, sensitivity tests and the original '
    '2014 source — is available on request.', italic=True)

doc.save(OUT)
print(f'Wrote {OUT}')
