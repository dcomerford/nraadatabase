"""Render point-by-point answers to Gerald's six follow-up questions as a Word .docx,
for Adrian to use in his reply. Same house style as build_gerald_response_docx.py
(Calibri, navy/grey palette, EEF3F8 header rows, FFF9E6 highlight, callout boxes).

Grounded in v5_calibration.py (the actual calibration) and the original 2014 source
at ozfclass.com/articles/mcsi/mcsi.html.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/dancomerford/Desktop/claude/nraadatabase/GRC_MCSI_Answers_to_Gerald_Questions.docx'

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
RED = RGBColor(0xB0, 0x2A, 0x2A)
HEADER_ROW_FILL = 'EEF3F8'
ROW_HIGHLIGHT = 'FFF9E6'
GREY_FILL = 'F4F4F4'
AMBER_BORDER = 'C9A227'
RED_BORDER = 'B02A2A'
GREEN_FILL = 'EAF3EA'
GREEN_BORDER = '3C7A3C'

doc = Document()
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _left_border(cell, hex_color, sz='24'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    borders.append(left)
    tcPr.append(borders)


def _add_runs(paragraph, text):
    text = text.replace('&amp;', '&')
    pattern = re.compile(r'(<b>.*?</b>|<i>.*?</i>)', re.S)
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith('<b>') and part.endswith('</b>'):
            paragraph.add_run(part[3:-4]).bold = True
        elif part.startswith('<i>') and part.endswith('</i>'):
            paragraph.add_run(part[3:-4]).italic = True
        else:
            paragraph.add_run(part)


def add_para(text, *, size=10.5, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if color is not None:
            r.font.color.rgb = color
    return p


def add_title(text):
    add_para(text, size=20, bold=True, color=NAVY, space_after=4)


def add_subtitle(text):
    add_para(text, size=12, color=GREY, space_after=6)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(15)
        r.bold = True
        r.font.color.rgb = NAVY


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = GREY


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10.5)


def add_small(text, italic=True):
    add_para(text, size=9, italic=italic, color=LIGHT_GREY, space_after=8)


def add_hr():
    p = doc.add_paragraph()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def add_callout(text, fill=HEADER_ROW_FILL, border=NAVY):
    border_hex = border if isinstance(border, str) else '1F4E78'
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(17)
    _shade_cell(cell, fill)
    _left_border(cell, border_hex)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(data, col_widths_cm=None, highlight_rows=None):
    n_rows, n_cols = len(data), len(data[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    table.style = 'Table Grid'
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for r_idx, row_data in enumerate(data):
        is_header = r_idx == 0
        is_highlight = highlight_rows and r_idx in highlight_rows
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if col_widths_cm and c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_runs(p, str(text))
            for r in p.runs:
                r.font.size = Pt(9.5)
                if is_header:
                    r.bold = True
                    r.font.color.rgb = NAVY
            if is_header:
                _shade_cell(cell, HEADER_ROW_FILL)
            elif is_highlight:
                _shade_cell(cell, ROW_HIGHLIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_gerald(text):
    """Gerald's question, quoted in a grey box."""
    add_callout('<i>Gerald asked:</i>  ' + text, fill=GREY_FILL, border='888888')


def add_send(text):
    """Ready-to-send answer, green left border."""
    add_callout('<b>Suggested reply →</b>  ' + text, fill=GREEN_FILL, border=GREEN_BORDER)


def add_action(text):
    """Action / caveat for Adrian before sending, amber."""
    add_callout('<b>⚑ Before you send:</b>  ' + text, fill=ROW_HIGHLIGHT, border=AMBER_BORDER)


# ── Header ───────────────────────────────────────────────────────────────────
add_title('MCSI V5 — Answers to Gerald’s Follow-up Questions')
add_subtitle('Prepared for Adrian · working notes for the reply to Gerald · July 2026')
add_hr()
add_para(
    'Adrian — Gerald’s six questions are answered below, each grounded in the actual '
    'calibration and, where relevant, the original published MCSI. For each one there is a '
    '<b>quote of his question</b>, a <b>suggested reply you can lift straight into your email</b>, '
    'and — where a claim needs your own confirmation before it goes out — a short '
    '<b>“before you send”</b> note. Two of these (the 2014 provenance and the acronym) need a '
    'quick check with Peter first; everything else is ready to send.')

# ── Provenance find ────────────────────────────────────────────────────────────
add_h2('Key discovery — we can now cite the original MCSI')
add_para(
    'The original MCSI is published online: <b>ozfclass.com/articles/mcsi/mcsi.html</b>. This '
    'lets us answer Gerald’s Question 1 with a source instead of an assertion. The important '
    'facts from it:')
add_bullet('<b>MCSI = “Mixed Category Score Index.”</b> (Not “Multiple” — see the note under '
           'Question 1.)')
add_bullet('<b>Built on 2014 Queen’s data:</b> <i>“Score data was collected from all Queens '
           'shoots (not leadups) held around Australia during calendar 2014.”</i> That confirms '
           'the 2014 / Queen’s provenance.')
add_bullet('<b>Sporter was not in it</b> — no mention anywhere, which backs our claim that '
           'Sporter was never part of the original build.')
add_bullet('<b>The original was intended for “low level team competition, and for miscellaneous '
           'award purposes… not… as a replacement for scoring of ranges and aggregates in '
           'individual events.”</b> That is exactly our Section 8 framing — a constructed award '
           'index, not a claim about the objectively best shooter.')
add_para('But it also shows three things V5 does <b>differently</b> from the 2014 original, and '
         'we should be honest about them rather than claim an unbroken lineage:')
add_bullet('<b>Different method.</b> The 2014 system derived its factors from <b>standard '
           'deviations of group size in minutes of angle</b> (SD ratios F-Std:TR 0.725, '
           'F-Open:TR 0.533, F/TR:TR 0.623). V5 uses <b>top-40% cohort averages</b>. These are '
           'different derivations — V5 is not a numerical continuation of the 2014 numbers.')
add_bullet('<b>Different centre weight.</b> The 2014 system added centres/super-Vs at '
           '<b>full value (weight 1.0)</b>. GRC uses <b>0.7</b> — so the 0.7 is a GRC choice, '
           'not inherited (this is Gerald’s Question 2).')
add_bullet('<b>Different name in circulation.</b> The original and Peter’s club note both say '
           '<b>“Mixed”</b>; the V5 response document says <b>“Multiple.”</b>')
add_callout(
    '<b>Bottom line for the reply:</b> position V5 as keeping the <b>purpose, scale and spirit</b> '
    'of the 2014 Mixed Category Score Index, but replacing its MOA method with a transparent, '
    'reproducible cohort calibration — and adding a properly-calibrated Sporter factor for the '
    'first time. That is fully defensible and turns every difference into a transparency argument.')

# ── Q1 ─────────────────────────────────────────────────────────────────────────
add_h2('1. Provenance of the original formula')
add_gerald(
    '“The original MCSI was built on King’s/Queen’s competition data from 2014 and earlier, and '
    'the club has used that formula for the last four years. You better be 100% certain of the '
    'accuracy of this statement! Explain where you got this data.”')
add_send(
    'The Mixed Category Score Index originates from a 2014 project that calibrated the Australian '
    'fullbore classes against one another using all of that year’s Queen’s shoots across the '
    'country — it is published at ozfclass.com/articles/mcsi/mcsi.html. Sporter was not part of '
    'it. GRC adopted the MCSI concept and has run a version of it for the last four years. V5 '
    'keeps that same purpose and scoring scale, but re-calibrates using recent (2024-onward) '
    'King’s/Queen’s data and a transparent cohort method, and calibrates Sporter properly for '
    'the first time.')
add_action(
    'Two accuracy points to lock down before this goes out. (a) The original data is '
    '<b>calendar 2014 Queen’s</b> specifically — the phrase “2014 <i>and earlier</i>” overstates '
    'it; drop “and earlier.” (b) Confirm with Peter <b>what GRC actually ran for four years</b> — '
    'Peter’s own club note describes a “median of the top half of shooters” method, which is '
    'neither the 2014 MOA method nor V5. Safest to say the club adopted the MCSI <i>concept</i> '
    'rather than claim it ran this exact formula, unless Peter confirms otherwise.')

# ── Q2 ─────────────────────────────────────────────────────────────────────────
add_h2('2. The 0.7 centre weight — and the brackets')
add_gerald(
    '“The 0.7 is still a black box number with no source quoted. Is it actually needed at all if '
    'you say the rankings are stable between 0.5 and 0.8? … the equation … should be bracketed: '
    'Adjusted MCSI = ( (score × conversion) + (centres × 0.7) ) × factor.”')
add_h3('On the brackets — he is right, adopt it')
add_send(
    'Agreed, and adopted. The bracketed form <b>( (Score × Conversion) + (Centres × 0.7) ) × '
    'Factor</b> is clearer and is mathematically identical to what we had, so no number changes.')
add_h3('On whether 0.7 is needed — yes; the value is a policy choice')
add_send(
    'The centre-weight term is needed: it is the only mechanism by which centres and V-bulls count '
    'toward the index, which is the whole premise that a well-centred near-possible can be a better '
    'shoot than a possible with few centres. What is <b>not</b> critical is the exact value — the '
    'rankings are stable anywhere from 0.5 to 0.8, so 0.7 is a robust policy setting rather than a '
    'fragile magic number. Its meaning is transparent: the break-even is 1 ÷ 0.7 = 1.43, i.e. it '
    'takes about 1½ extra centres to make up one dropped point. Note the original 2014 MCSI counted '
    'centres at full value (1.0); GRC deliberately moderated that to 0.7 so centres inform the '
    'ranking without dominating it.')
add_para(
    'On Gerald’s specific point that removing the 0.7 would require re-doing the factors — '
    '<b>he is exactly right</b>, and it is a good reason to keep the term rather than drop it. '
    'The factors are derived from each discipline’s cohort average of '
    '<i>(Score × Conversion + Centres × 0.7)</i>, and the common target (94.471) was itself '
    'computed with the 0.7 in place. Setting the weight to zero would (a) remove centres from the '
    'championship entirely and (b) force every factor to be re-derived. So the honest position is: '
    'the term stays; only its value is a tuning choice, and 0.7 sits safely inside the stable band.')

# ── Q3 ─────────────────────────────────────────────────────────────────────────
add_h2('3. Where 94.471 comes from')
add_gerald('“Where does the 94.471 come from? The calculation of this number needs to be '
           'explained.”')
add_send(
    '94.471 is the common target every discipline is lifted onto. It is the <b>average of the '
    'three F-class disciplines’ top-40% cohort scores measured under the existing factors</b> — '
    'i.e. the level F-class already sat at. We chose F-class as the anchor deliberately: it means '
    'the F-classes barely move, and only the mis-calibrated classes (Sporter, and drift in TR / '
    'F-Open) are pulled onto that shared level. The arithmetic:')
add_table([
    ['F-class discipline', 'Top-40% cohort average', '× existing factor', '= level'],
    ['F-Open', '67.198', '1.42', '95.42'],
    ['F-Standard', '64.062', '1.46', '93.53'],
    ['F/TR', '65.165', '1.45', '94.49'],
    ['<b>Average → target</b>', '', '', '<b>≈ 94.47</b>'],
], col_widths_cm=[5.5, 4.5, 3.5, 3.5], highlight_rows={4})
add_para(
    'Each discipline’s factor is then just <b>94.471 ÷ (its own top-40% cohort average)</b>. The '
    'units resolve because the ~94 target (after factor) is divided by the ~67 cohort average '
    '(before factor): 94.471 ÷ 66.907 = 1.412 for Target Rifle, and so on.')
add_action(
    'Be ready for the obvious follow-up — <i>“isn’t it circular to use the old factors to set the '
    'new target?”</i> Yes, and intentionally: F-class was considered about right, so we anchor to '
    'it and fix the outliers rather than re-normalise everything from scratch. Say that plainly if '
    'asked. Also note the response document’s phrase <i>“the average of the three F-class cohorts”</i> '
    'is too loose — it invites the reading that 94.471 is the average of the raw means (65.5). '
    'Reword it to the definition above.')

# ── Q4 ─────────────────────────────────────────────────────────────────────────
add_h2('4. “Top-40% cohort mean converted score” — and the top-60% question')
add_gerald(
    '“What do you mean by top-40% cohort mean converted score? Are you only taking the top 60% of '
    'all records, if so why?”')
add_send(
    'It is the top <b>40%</b> — the best 40% of scores — not the top 60%. For each combination of '
    'state, year, discipline and range we take the best 40% of that field (with a floor of 5 '
    'scores so small fields still count), then pool them for the discipline. The “converted score” '
    'is <i>(Score × Conversion + Centres × 0.7)</i>, and we take its average over that pooled top '
    '40%. We use the top 40% rather than the whole field because the full field is dominated by '
    'casual and beginner entrants whose spread varies between disciplines for reasons unrelated to '
    'the scoring scale; the championship is decided at the sharp end, so calibrating on competitive '
    'scores is far more stable. It is the same instinct as the top-half approach used in the club '
    'system, just tightened to 40%.')

# ── Q5a ────────────────────────────────────────────────────────────────────────
add_h2('5. What “cohort” means')
add_gerald('“What do you mean by the term ‘cohort’?”')
add_send(
    'A “cohort” is simply <b>the top 40% of scores in a discipline</b> — that discipline’s pool of '
    'competitive performances across all the King’s/Queen’s ranges in the dataset. We will drop the '
    'jargon in the report and write “the top 40% of scores in each discipline” instead, which is '
    'plainer.')

# ── Q5b ────────────────────────────────────────────────────────────────────────
add_h2('6. “Average” vs “mean” — consistent wording')
add_gerald(
    '“…the factor equalises each discipline’s cohort average. Your terminology should be '
    'consistent: use ‘average’ or ‘mean’ but don’t use both… I suggest ‘average.’”')
add_send(
    'Agreed — we will use <b>“average”</b> throughout and remove “mean.” The underlying calculation '
    'is an arithmetic average, so the plainer word is also the accurate one.')

# ── Naming + proposal ──────────────────────────────────────────────────────────
add_h2('Two more housekeeping items Gerald will notice')
add_h3('The acronym')
add_action(
    'The response document is titled <b>“Multiple</b> Category Score Index.” Both the original '
    'source and Peter’s club note call it the <b>“Mixed</b> Category Score Index.” In a document '
    'arguing for rigour this is exactly the kind of slip a reviewer flags — recommend standardising '
    'on <b>“Mixed”</b> everywhere. Worth a quick word with Peter to confirm the club’s preferred '
    'form.')
add_h3('The two-page committee proposal Gerald asked for')
add_para(
    'Gerald also asked for the <b>simple, under-two-page proposal</b> to circulate at the meeting, '
    'with the deeper detail held in reserve. The full technical response is the “deeper answer” '
    'reserve; the two-pager should be just: (1) what the MCSI is and is for, (2) the one formula in '
    'its bracketed form, (3) the factor table with record counts, (4) the three things V5 changes '
    '(recent data, transparent method, Sporter), and (5) the governance rule (factors locked per '
    'season). Happy to draft that next.')

add_hr()
add_para(
    'Everything marked <b>“Suggested reply”</b> is ready to send. The <b>amber</b> notes '
    '(2014 wording, the 94.471 circularity follow-up, and the Mixed/Multiple name) are the only '
    'places you need your own confirmation first — mostly a quick check with Peter.', italic=True)

doc.save(OUT)
print(f'Wrote {OUT}')
