"""Render the short Response to Gerald's V5 MCSI review as a Word .docx.

Mirrors GRC_MCSI_Response_to_Gerald.pdf and the committee-report docx style
(Calibri, navy/grey palette, EEF3F8 header rows, FFF9E6 highlight, callout boxes).
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/dancomerford/Desktop/claude/nraadatabase/GRC_MCSI_Response_to_Gerald.docx'

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
HEADER_ROW_FILL = 'EEF3F8'
ROW_HIGHLIGHT = 'FFF9E6'
GREY_FILL = 'F4F4F4'
AMBER_BORDER = 'C9A227'

doc = Document()
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _left_border(cell, hex_color, sz='24'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), hex_color)
    borders.append(left)
    tcPr.append(borders)


def _add_runs(paragraph, text):
    text = text.replace('&amp;', '&')
    pattern = re.compile(r'(<b>.*?</b>|<i>.*?</i>)', re.S)
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith('<b>') and part.endswith('</b>'):
            paragraph.add_run(part[3:-4]).bold = True
        elif part.startswith('<i>') and part.endswith('</i>'):
            paragraph.add_run(part[3:-4]).italic = True
        else:
            paragraph.add_run(part)


def add_para(text, *, size=10.5, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if color is not None:
            r.font.color.rgb = color
    return p


def add_title(text):
    add_para(text, size=20, bold=True, color=NAVY, space_after=4)


def add_subtitle(text):
    add_para(text, size=12, color=GREY, space_after=6)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(15)
        r.bold = True
        r.font.color.rgb = NAVY


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10.5)


def add_small(text, italic=True):
    add_para(text, size=9, italic=italic, color=LIGHT_GREY, space_after=8)


def add_hr():
    p = doc.add_paragraph()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def add_callout(text, fill=HEADER_ROW_FILL, border=NAVY):
    border_hex = border if isinstance(border, str) else '1F4E78'
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(17)
    _shade_cell(cell, fill)
    _left_border(cell, border_hex)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(data, col_widths_cm=None, highlight_rows=None):
    n_rows, n_cols = len(data), len(data[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False
    table.style = 'Table Grid'
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for r_idx, row_data in enumerate(data):
        is_header = r_idx == 0
        is_highlight = highlight_rows and r_idx in highlight_rows
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if col_widths_cm and c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_runs(p, str(text))
            for r in p.runs:
                r.font.size = Pt(9.5)
                if is_header:
                    r.bold = True
                    r.font.color.rgb = NAVY
            if is_header:
                _shade_cell(cell, HEADER_ROW_FILL)
            elif is_highlight:
                _shade_cell(cell, ROW_HIGHLIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Header ───────────────────────────────────────────────────────────────────
add_title("Technical Response — Review of the V5 MCSI")
add_subtitle('Geelong Rifle Club — Multiple Category Score Index, Version 5 · '
             'Prepared for the GRC committee · July 2026')
add_hr()
add_para(
    'This document responds to the technical review of the V5 MCSI. Most of the points raised '
    'are fair, several improve the model, and a few rest on a misunderstanding of what the MCSI '
    'is for. Each is answered directly below, with the actual numbers from our dataset.')

# ── 1 ────────────────────────────────────────────────────────────────────────
add_h2('1. What the MCSI is — and what it is not')
add_para(
    'This is the most important point, because several objections dissolve once it is clear. '
    'The MCSI was <b>not</b> built as a club-level tool and it does <b>not</b> compare GRC '
    'members against elite shooters as individuals. The original MCSI was built on '
    'King’s/Queen’s competition data from <b>2014 and earlier</b>, and the club has used that '
    'formula for the last four years. The current V5 work is an <b>update</b> of it using '
    'recent competition data.')
add_callout(
    '<b>The large external dataset is used to estimate the shape and relative behaviour of '
    'each discipline’s score distribution — not to create individual shooter ratings.</b>')
add_para(
    'No GRC member is being told their score equals a national champion’s. The model asks '
    '<i>“how are scores distributed within each discipline — how compressed, how often is a '
    'possible reached, how do the top performers cluster?”</i> It does <b>not</b> ask '
    '<i>“is this member as skilled as a King’s medallist?”</i> Because King’s events produce '
    'genuinely large, statistically significant samples in each discipline, we can measure '
    'those distributions reliably — which is exactly why that data is used.')
add_para('Its purpose is narrow and worth stating plainly:')
add_callout(
    'To convert performances recorded under different scoring systems (50-point vs 60-point) '
    'and equipment conditions into a <b>common index</b> that gives a reasonably equitable '
    'basis for comparing <b>relative</b> performance across GRC disciplines.')
add_para(
    '<b>Two things have changed since the original formula, and they are the reason for the V5 '
    'update.</b> First, <b>Sporter was never part of the 2014 build.</b> For the last few years '
    'it has been carried in the championship using an <i>estimated</i> factor that was, in '
    'effect, a placeholder rather than a calibrated value. We now have enough recent '
    'King’s/Queen’s data to calibrate Sporter properly — and we have demonstrated that the '
    'sample is large enough to do so (Section 3). Second, <b>equipment has moved on</b> — most '
    'visibly in F-Open — so the older factors no longer reflect current scoring.')
add_para(
    'The V5 update therefore uses the recent dataset to produce factors representative of the '
    'present era. One honest limitation: <b>Sporter Open and Sporter (Production) were combined '
    'in the earlier records</b>, so at this stage they cannot be cleanly separated — V5 uses a '
    'single merged Sporter factor. That can be revisited once the two are recorded distinctly '
    'for long enough to calibrate separately.')
add_para(
    'None of this re-normalises the formula on GRC’s own small sample. The aim is to (a) properly '
    'include the smaller classes — especially Sporter — which were previously under-represented '
    'or estimated, and (b) rebalance the disciplines so no single class holds a structural '
    'advantage. Celebrating individual discipline champions in their own right remains entirely '
    'compatible with this — the two are not mutually exclusive.')

# ── 2 ────────────────────────────────────────────────────────────────────────
add_h2('2. “The factors are a black box / magic numbers” — they are not')
add_para(
    'This is a fair criticism of the <i>report</i>, not the <i>calculator</i>. The derivation '
    'was never published; here it is. The V5 formula is:')
add_callout('<b>Adjusted MCSI = (Score × Conversion + Centres × 0.7) × Factor</b>',
            fill=GREY_FILL, border='444444')
add_para(
    'Each discipline’s factor is <b>not chosen by eye</b>. It is the single number that lifts '
    'that discipline’s top-40% cohort average onto one common target (94.471, the average of '
    'the three F-class cohorts):')
add_callout('<b>Factor = 94.471 ÷ (that discipline’s top-40% cohort mean converted score)</b>')
add_table([
    ['Discipline', 'Top-40% cohort mean', '94.471 ÷ mean', 'V5 factor'],
    ['Target Rifle', '66.907', '1.412', '<b>1.412</b>'],
    ['F-Open', '67.198', '1.406', '<b>1.406</b>'],
    ['F-Standard', '64.062', '1.475', '<b>1.475</b>'],
    ['F/TR', '65.165', '1.450', '<b>1.450</b>'],
    ['Sporter (Open + Prod. merged)', '68.292', '1.383', '<b>1.383</b>'],
], col_widths_cm=[7.0, 4.0, 3.0, 3.0])
add_para(
    'The factor is fully reproducible — it can be recomputed independently from the cohort '
    'means. The discipline whose strong scores average <b>lowest</b> (F-Standard, 64.06) gets the '
    '<b>largest</b> factor; the one that averages <b>highest</b> (Sporter, 68.29) gets the '
    '<b>smallest</b>. There is nothing magic in it: it is one division per discipline. '
    '<b>The final report will publish this table with the underlying record counts.</b>')

# ── 3 ────────────────────────────────────────────────────────────────────────
add_h2('3. “Why King’s/Queen’s data, not club data?” — sample size')
add_para(
    'The instinct to use our own club data is reasonable, but the numbers make it impossible '
    'today. The V5 calibration uses <b>~19,800 individual range scores across 135 championship '
    'days</b>:')
add_table([
    ['Discipline', 'Scores', 'Distinct shooters'],
    ['Target Rifle', '7,510', '351'],
    ['F-Open', '4,082', '222'],
    ['F-Standard', '3,731', '193'],
    ['F/TR', '2,437', '124'],
    ['Sporter', '2,018', '152'],
], col_widths_cm=[9.0, 4.0, 4.0])
add_para(
    'A stable calibration needs roughly 3,000–4,000 scores <b>per discipline</b>. Geelong '
    'produces a small fraction of that in a season. A GRC-only factor wouldn’t describe '
    '<i>the discipline</i> — it would describe <b>the two or three regulars who happened to '
    'turn out</b>, and it would lurch year to year. We would need to roughly <b>5× our '
    'participation</b> before a club-only model became statistically meaningful — and for '
    'Sporter, not even then.')
add_para(
    '<b>We have also tested this directly.</b> We applied the V5 factors to the most recent New '
    'South Wales King’s data — <b>data not used to derive them</b> — and the factors held: the '
    'results were consistent. When that new block of championship data was folded into the '
    'existing dataset, the factors <b>barely moved</b> (a marginal shift on F-Standard only). '
    'This matters: if our sample were too small to be statistically meaningful, adding a fresh '
    'season of championship data would have swung the factors noticeably. It did not — which is '
    'direct evidence the sample is large enough to rely on.')
add_para(
    'This is exactly why the external data is used for the <i>distribution shape</i> while GRC '
    'data is the <i>local sanity check</i>. Elite data isn’t a perfect mirror of club shooting; '
    'the honest long-term answer is a <b>hybrid</b> — anchor on the national distribution now, '
    'and give GRC data progressively more weight as it accumulates (credibility weighting).')

# ── 4 ────────────────────────────────────────────────────────────────────────
add_h2('4. “All-else-equal, F-Standard always wins” — correct, and here’s why')
add_para(
    'This is a valid observation. When every discipline is put on the same percentage of '
    'possible with the same centres, F-Standard leads every row. Reproducing the test exactly '
    '(perfect score, 10 centres):')
add_table([
    ['% of possible', 'TR', 'F-Open', 'F-Std', 'FTR', 'Sporter', 'Wins'],
    ['100%', '94.60', '94.20', '<b>98.83</b>', '97.15', '92.66', 'F-Std'],
    ['98%',  '92.91', '92.51', '<b>97.06</b>', '95.41', '91.00', 'F-Std'],
    ['96%',  '91.22', '90.83', '<b>95.28</b>', '93.67', '89.34', 'F-Std'],
    ['90%',  '86.13', '85.77', '<b>89.98</b>', '88.45', '84.36', 'F-Std'],
], col_widths_cm=[2.8, 1.9, 2.2, 2.2, 1.9, 2.2, 1.8])
add_small('Order: F-Std → FTR → TR → F-Open → Sporter.')
add_para(
    '<b>Why it happens:</b> the factor equalises each discipline’s <i>cohort average</i> '
    '(Section 2), not its <i>maximum</i>. F-Standard’s top shooters sit furthest below their '
    'ceiling, so F-Std earns the biggest multiplier. When you then feed in a score at or near '
    'the <b>maximum</b> — i.e. above F-Std’s own cohort average — that large multiplier '
    'over-rewards it. It is a mathematical property of <b>mean-equalisation</b>, not a hidden '
    'preference for F-Standard.')
add_para('This must be acknowledged as a real trade-off. Two honest fixes, if the committee '
         'finds it unacceptable:')
add_bullet('<b>Factor shrinkage</b> — pull all factors part-way toward a common value, keeping '
           'half the calibration but reducing the equal-score spread; or')
add_bullet('calibrate to a <b>top-cohort / possibles</b> target instead of the mean (this '
           'equalises the elite tail rather than the average — a different, defensible choice).')
add_callout(
    '<b>NOTE — factor consistency:</b> an earlier draft circulated an F-Standard factor of '
    '<b>1.440</b>; the current calibration uses <b>1.475</b>. At 1.475, F-Std has the largest '
    'factor — which is exactly the behaviour described above. Every figure in this document uses '
    '<b>1.475</b>; the final report must too.',
    fill=ROW_HIGHLIGHT, border=AMBER_BORDER)

# ── 5 ────────────────────────────────────────────────────────────────────────
add_h2('5. Does a possible always beat a near-possible? — no, but only a 2-centre swing '
       'flips it')
add_para(
    'This is also correct, and the earlier claim that “a possible always beats a near-possible” '
    'should be withdrawn. Within a single discipline, ranking is by (score + 0.7 × centres) — '
    'the factor cancels, so the comparison is identical in every class. A dropped point is worth '
    '<b>1</b>; a centre is worth <b>0.7</b>. The break-even is 1 ÷ 0.7 = <b>1.43 centres</b> — '
    'so a near-possible must carry at least <b>2 more centres</b> than the possible to overtake '
    'it:')
add_table([
    ['Possible', 'Near-possible', 'score + 0.7 × centres', 'Winner'],
    ['50.8', '49.9', '55.6 vs 55.3', '<b>50 holds</b>'],
    ['50.7', '49.9', '54.9 vs 55.3', '<b>49 wins</b>'],
    ['50.10', '49.10', '57.0 vs 56.0', '<b>50 wins</b>'],
], col_widths_cm=[3.0, 3.5, 5.5, 5.0])
add_small('Shown for a 50-max discipline; the same 2-centre rule applies in F-class — 60 vs 59.')
add_para('So a possible with a normal centre count is <b>robust</b>: only a near-possible '
         'carrying two extra centres — an uncommon result — overturns it. Whether even that '
         'should be allowed is a <b>policy choice</b>, not a maths error:')
add_bullet('<b>Policy 1 (primary absolute):</b> a higher score always wins; centres only break '
           'ties.')
add_bullet('<b>Policy 2 (combined performance — current V5):</b> a large enough centre '
           'advantage can outweigh one dropped point.')
add_para(
    'The committee should pick one deliberately. The report will be corrected to say: '
    '<i>the system strongly rewards primary score and generally favours possibles, but does '
    'not guarantee that every possible outranks every near-possible.</i>')

# ── 6 ────────────────────────────────────────────────────────────────────────
add_h2('6. Distance — correct in principle, but deliberately not adopted')
add_para(
    'It is correct that a 50.10 at 900m is, on average, harder than at 500m, and V5 uses one '
    'factor per discipline across all distances. <b>We have looked at this.</b> Two reasons it '
    'is not in V5:')
add_bullet(
    '<b>Data.</b> A distance-specific factor means <b>6 disciplines × 5 distances = 30 cells</b>, '
    'each needing ~100+ scores from 15+ shooters to be stable. We have that depth for some '
    'disciplines but not others — <b>Sporter in particular</b> — and there is a real risk of '
    '<b>double-counting difficulty</b> that is already reflected in the observed score '
    'distribution.')
add_bullet(
    '<b>Usability.</b> A separate factor for every distance, in both <b>metres and yards</b>, '
    'across every discipline, would turn the formula into exactly the “black box” the review '
    'warns against. We want the MCSI to stay <b>usable — calculable by any shooter by hand.</b> '
    'One factor per discipline keeps it transparent.')
add_para(
    'This is a legitimate <b>future (V6) development item</b>, subject to sample-size thresholds — '
    'not a V5 defect.')

# ── 7 ────────────────────────────────────────────────────────────────────────
add_h2('7. “Factors should evolve” — agreed, under governance')
add_para(
    'Yes. Equipment (barrels, projectiles, F-Open cartridge design) has changed markedly in a '
    'decade, and the factors should track it. But they must <b>not</b> drift continuously, or '
    'a January result wouldn’t compare with a November one. The process should be: <b>lock '
    'factors before each season → collect scores → annual post-season review → test '
    'proposed changes against history → committee approval → apply from the next season '
    'only.</b> Every factor set versioned, historical results never rewritten. We’ll add this '
    'as a formal governance section.')

# ── 8 ────────────────────────────────────────────────────────────────────────
add_h2('8. “Is the MCSI goal even possible?”')
add_para(
    'In a strict scientific sense, there is no single “best shooter” across disciplines that '
    'use different equipment and skills — that is philosophically correct. But the MCSI '
    'never claimed that. It is a <b>constructed index</b>, in exactly the same family as golf '
    'handicaps, the decathlon scoring tables, and motorsport balance-of-performance. None of '
    'those prove two performances are physically identical; they provide an agreed, transparent '
    'framework for comparison. The right description of the title is:')
add_callout(
    '<b>GRC Multiple Category Champion under the adopted MCSI rules</b> — not “the objectively '
    'best shooter.”')
add_para(
    'Four years of work shouldn’t be discarded over the apples-to-oranges objection; and '
    'separately celebrating each discipline’s own champion costs us nothing.')

# ── 9 ────────────────────────────────────────────────────────────────────────
add_h2('9. What we will change in the report')
for i, item in enumerate([
    '<b>Publish the factor derivation</b> — the table in Section 2, with record counts, so the '
    'factors are reproducible, not “magic.”',
    '<b>Fix the F-Standard factor inconsistency</b> — use <b>1.475</b> throughout.',
    '<b>Add a centre-weight sensitivity test</b> — show rankings are stable across w = 0.5–0.8; '
    'describe 0.7 as a calibrated policy value, not a universal constant.',
    '<b>Correct the possible/near-possible wording</b> (Section 5) and state the chosen policy '
    'explicitly.',
    '<b>Acknowledge the equal-score F-Std lead</b> (Section 4) and offer factor-shrinkage as '
    'the mitigation if the committee wants it.',
    '<b>Add a distance-analysis section</b> recording why per-distance factors are deferred to '
    'V6 (data gaps + usability), subject to sample-size thresholds.',
    '<b>Add a factor-governance section</b> — annual review, seasonal locking, versioning, no '
    'retrospective changes.',
], 1):
    add_bullet(f'{i}. {item}')
add_hr()
add_para(
    'The review doesn’t show V5 is wrong — it shows exactly where the <i>explanation</i> '
    'needs strengthening before we present it as final. We’re happy to walk through any of the '
    'above.')

doc.save(OUT)
print(f'Wrote {OUT}')
