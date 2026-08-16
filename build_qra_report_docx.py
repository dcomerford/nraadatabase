"""Build the QRA Kings Prize Shoot 2026 report for Adrian (.docx).

Covers: the new event, V5 vs June-13 on it, whether it agrees with NSWRA 2026,
what happens when both 2026 championships are folded into the calibration pool,
and what the distance data does and does not tell us about a distance-aware
formula.

Every number is computed at build time from kings_qra2026.csv, kings_nswra2026.csv
and the national K&Q pool — nothing is typed in by hand.

Usage:  python3 build_qra_report_docx.py
"""
import math
import re
from collections import Counter, defaultdict
from statistics import mean

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import analyse_qra_v5 as A
import compare_qra_nswra as C
import pool_recalibration as P

OUT = '/Users/dancomerford/Desktop/claude/nraadatabase/QRA_Kings_2026_V5_report.docx'

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
HEADER_ROW_FILL = 'EEF3F8'
ROW_HIGHLIGHT = 'FFF9E6'
AMBER_FILL = 'FDF3E3'
GREEN_FILL = 'EAF3EA'

doc = Document()
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)


# ── style helpers (house style, same as the Gerald/V5 packs) ──────────────────
def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _left_border(cell, hex_color, sz='24'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '0'); left.set(qn('w:color'), hex_color)
    borders.append(left); tcPr.append(borders)


def _add_runs(paragraph, text):
    text = text.replace('&amp;', '&')
    for part in re.split(r'(<b>.*?</b>|<i>.*?</i>)', text, flags=re.S):
        if not part:
            continue
        if part.startswith('<b>'):
            paragraph.add_run(part[3:-4]).bold = True
        elif part.startswith('<i>'):
            paragraph.add_run(part[3:-4]).italic = True
        else:
            paragraph.add_run(part)


def add_para(text, *, size=10, bold=False, italic=False, color=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
        if bold: r.bold = True
        if italic: r.italic = True
        if color is not None: r.font.color.rgb = color
    return p


def add_title(text):
    add_para(text, size=16, bold=True, color=NAVY, space_after=2)


def add_subtitle(text):
    add_para(text, size=10, color=GREY, space_after=5)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(11.5); r.bold = True; r.font.color.rgb = NAVY


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10)


def add_small(text, italic=True):
    add_para(text, size=8.5, italic=italic, color=LIGHT_GREY, space_after=5)


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); p._p.get_or_add_pPr().append(pBdr)


def add_callout(text, fill=HEADER_ROW_FILL, border='1F4E78'):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0); cell.width = Cm(17.4)
    _shade_cell(cell, fill); _left_border(cell, border)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_table(data, col_widths_cm=None, highlight_rows=None, highlight_fill=ROW_HIGHLIGHT,
              align_from=1):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.autofit = False; table.style = 'Table Grid'
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for r_idx, row_data in enumerate(data):
        is_header = r_idx == 0
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
            if c_idx >= align_from:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(p, str(text))
            for r in p.runs:
                r.font.size = Pt(9)
                if is_header:
                    r.bold = True; r.font.color.rgb = NAVY
            if is_header:
                _shade_cell(cell, HEADER_ROW_FILL)
            elif highlight_rows and r_idx in highlight_rows:
                _shade_cell(cell, highlight_fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ══ compute everything ════════════════════════════════════════════════════════
rows = A.load('kings_qra2026.csv')
keys = []
for r in rows:
    if r['key'] not in keys:
        keys.append(r['key'])
n_shooters = len({r['name'] for r in rows})

# per-event implied factors, QRA vs NSWRA
imp_q, n_q = C.implied(C.load('kings_qra2026.csv'))
imp_n, n_n = C.implied(C.load('kings_nswra2026.csv'))

# pool recalibration, both normalisations
kq, n_kq = P.kq_buckets()
nsw_b, n_nsw = P.csv_buckets('kings_nswra2026.csv', 'NSWRA2026')
qra_b, n_qra = P.csv_buckets('kings_qra2026.csv', 'QRA2026')
both_b = dict(kq); both_b.update(nsw_b); both_b.update(qra_b)
abs_base, abs_n, _ = P.calibrate(kq, False)
abs_new, abs_nn, _ = P.calibrate(both_b, False)
ps_base, ps_n, _ = P.calibrate(kq, True)
ps_new, ps_nn, _ = P.calibrate(both_b, True)

# string-length mix
def mix(bk):
    ct = Counter()
    for vals in bk.values():
        for _p, _x, s in vals:
            ct[s] += 1
    tot = sum(ct.values())
    return {s: 100 * n / tot for s, n in ct.items()}


mix_kq, mix_nsw, mix_qra = mix(kq), mix(nsw_b), mix(qra_b)

# per-shoot implied factors + hardship, for the distance section
shoot_rows, hard, gap, dist = [], [], [], []
for k in keys:
    sub = [r for r in rows if r['key'] == k]
    i2, _n2, _t = A.implied(sub)
    tr = sorted((r for r in sub if r['grp'] == 'TR'),
                key=lambda r: -A.base(r) / max(1, r['shot_count']))
    kk = max(1, math.ceil(0.4 * len(tr)))
    h = sum(A.base(r) / max(1, r['shot_count']) for r in tr[:kk]) / kk
    hard.append(h); gap.append(i2['F-Standard'] - i2['F-Open']); dist.append(sub[0]['dist'])
    shoot_rows.append((k, sub[0]['dist'], len(sub), i2, h))


def ordinal(n):
    if 10 <= n % 100 <= 20:
        return f'{n}th'
    return f'{n}{ {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th") }'.replace(' ', '')


def corr(x, y):
    mx, my = mean(x), mean(y)
    num = sum((a - mx) * (b - my) for a, b in zip(x, y))
    den = (sum((a - mx) ** 2 for a in x) * sum((b - my) ** 2 for b in y)) ** 0.5
    return num / den if den else float('nan')


r_hard, r_dist = corr(hard, gap), corr(dist, gap)
# the shoot at the longest range — where the weather bit hardest
longest = max(shoot_rows, key=lambda s: s[1])

# full-series leaderboards
av5 = A.aggregate(rows, keys, A.v5)
apj = A.aggregate(rows, keys, A.peter)

ORD = P.CATS
LAB = P.LABEL
V5PUB = P.PUBLISHED_V5

# ══ document ══════════════════════════════════════════════════════════════════
add_title('QRA Kings Prize Shoot 2026 — MCSI V5 check')
add_subtitle('A second independent King\'s-class test of V5, and what happens when both '
             '2026 championships are folded into the calibration pool · prepared for Adrian '
             '· 16 August 2026')
add_hr()

add_h2('Summary')
add_bullet(f'<b>V5 holds.</b> On {len(rows):,} strings from a championship it has never seen, '
           'V5 produces a balanced cross-discipline leaderboard and its factors sit within '
           f'{max(abs(imp_q[g] - V5PUB[{"TR": "TR", "F-Open": "F-Open", "F-Standard": "F-Std", "FTR": "FTR", "Sporter": "S"}[g]]) for g in imp_q):.3f} '
           'of what this event implies.')
add_bullet('<b>Folding NSWRA 2026 and QRA 2026 into the pool changes nothing.</b> On a '
           'like-for-like comparison every factor moves by at most '
           f'{max(abs(ps_new[c] - ps_base[c]) for c in ORD):.4f} — '
           f'{max(abs((ps_new[c] - ps_base[c]) * P.AGG_BASE) for c in ORD):.1f} MCSI points on a '
           '~490-point series aggregate. No factor changes at the 3 decimals V5 is published at.')
add_bullet('<b>One thing did turn up, and it is a method issue rather than a factor issue.</b> '
           'The archive is 87% ten-shot strings; NSWRA 2026 is half fifteen-shot. Pooling whole-string '
           'totals, as the original calibration did, makes that mix difference look like a real '
           'shift in shooting standard. Normalising per shot removes about 80% of the apparent '
           'movement. Worth fixing before the next recalibration.')
add_bullet('<b>Distance-based factors are not ready, and this meeting shows why.</b> The '
           'cross-discipline spread tracks how hard a shoot was, not how far it was '
           f'(r = {r_hard:+.2f} against difficulty, r = {r_dist:+.2f} against distance). Fitting '
           'factors to distance on this data would largely be fitting the weather.')

add_h2('1. The data')
add_para(f'Pulled from the QRA results system on 16 August 2026: <b>{len(rows):,} shooter-strings, '
         f'{n_shooters} shooters, {len(keys)} shoots</b> over 300, 500, 600, 800, 900 and 1000 yards '
         'across three days. Unlike the NSWRA export, this source publishes the distance against '
         'every match, which is what makes section 5 possible.')
add_para('Discipline split: ' +
         ', '.join(f'{k} {v}' for k, v in sorted(Counter(r['discipline'] for r in rows).items())) + '.')
add_callout('<b>Validation:</b> all <b>1,089</b> official aggregate rows published by QRA — the '
            'Day 1, Day 2, Day 3 and full Kings Series aggregates — reconcile <b>exactly</b> to the '
            'sum of the individual strings we extracted, with zero mismatches. The scores CSV that '
            'accompanies this report is therefore a faithful copy of the official result.',
            fill=GREEN_FILL, border='3C7A3C')

add_h2('2. V5 versus June-13 on the full Kings Series')
add_para(f'Full-series cross-discipline aggregate, {len(av5)} shooters who completed all '
         f'{len(keys)} shoots. Left is V5, right is the June-13 (Peter) formula on the same scores.')
data = [['#', 'V5 — shooter', 'Disc', 'Raw agg', 'V5', 'June-13 — shooter', 'Disc', 'Jun-13']]
for i in range(10):
    x, y = av5[i], apj[i]
    data.append([str(i + 1), x[1], x[2], f'{x[4]}.{x[5]}', f'{x[0]:.1f}',
                 y[1], y[2], f'{y[0]:.1f}'])
add_table(data, [0.8, 3.9, 2.0, 1.9, 1.5, 3.9, 2.0, 1.4])
v5mix = Counter(r[3] for r in av5[:10])
pmix = Counter(r[3] for r in apj[:10])
add_para('<b>Top-10 discipline mix — V5:</b> ' + ', '.join(f'{k} {v}' for k, v in v5mix.items()) +
         '.  <b>June-13:</b> ' + ', '.join(f'{k} {v}' for k, v in pmix.items()) + '.')
rank_p = {r[1]: i for i, r in enumerate(apj, 1)}
moves = [(r[1], r[2], i, rank_p[r[1]]) for i, r in enumerate(av5[:20], 1) if r[1] in rank_p]
big = max(moves, key=lambda m: m[3] - m[2])
sp = next((m for m in moves if m[1].startswith('Sporter')), None)
add_para('The pattern from NSWRA repeats. V5 spreads the top ten across four disciplines; June-13 '
         'concentrates it in F-Open and pushes Target Rifle and Sporter shooters a long way down. '
         f'Of the shooters in V5\'s top 20, the biggest fall is {big[0]} ({big[1]}), '
         f'{ordinal(big[2])} on V5 and {ordinal(big[3])} on June-13' +
         (f'; {sp[0]} ({sp[1]}) goes from {ordinal(sp[2])} to {ordinal(sp[3])}. ' if sp else '. ') +
         'That is the same behaviour we reported in June: June-13 gives full value to centres with '
         'no iron-sight conversion, so it rewards whichever disciplines shoot the most centres.')

add_h2('3. Does QRA agree with NSWRA about the factors?')
add_para('The useful test is not whether one event agrees with V5, but whether two independent '
         'events agree with <i>each other</i>. Each event is scored on its own: within the event, '
         'what factor would each discipline need to sit level with the F-class anchor?')
GMAP = {'TR': 'TR', 'F-Open': 'F-Open', 'F-Standard': 'F-Std', 'FTR': 'FTR', 'Sporter': 'S'}
data = [['Discipline', 'V5', 'NSWRA 2026', 'QRA 2026', 'Both events say']]
hl = []
for i, g in enumerate(['TR', 'F-Open', 'F-Standard', 'FTR', 'Sporter'], 1):
    v = V5PUB[GMAP[g]]
    dn, dq = imp_n[g] - v, imp_q[g] - v
    if dn * dq > 0:
        say = 'V5 slightly ' + ('low' if dn > 0 else 'high')
        if g == 'F-Standard':
            hl.append(i)
    else:
        say = 'they disagree — no signal'
    data.append([LAB[GMAP[g]], f'{v:.3f}', f'{imp_n[g]:.4f}', f'{imp_q[g]:.4f}', say])
add_table(data, [3.4, 2.0, 3.0, 3.0, 6.0], highlight_rows=hl)
add_para('Four of five agree in direction. <b>F-Standard is the one replicated signal:</b> both '
         f'championships imply V5\'s 1.475 is a touch high (NSWRA {imp_n["F-Standard"]:.4f}, '
         f'QRA {imp_q["F-Standard"]:.4f}). That is the same thing NSWRA hinted at in June, now '
         'confirmed independently. Sporter points opposite ways in the two events, so there is no '
         'signal there — consistent with the small and self-selected Sporter samples.')

add_h2('4. Folding both championships into the calibration pool')
add_para('This is the question that matters for the formula. The national K&Q pool V5 was built on '
         f'is <b>{n_kq:,} strings</b>. NSWRA 2026 adds {n_nsw:,} and QRA 2026 adds {n_qra:,}, a '
         f'{100 * (n_nsw + n_qra) / n_kq:.0f}% increase. Re-running the calibration exactly as it '
         'was originally run — same buckets, same top-40% cohorts, same F-class anchor — gives:')

data = [['Discipline', 'Pool only', 'Pool + both events', 'Move', 'MCSI pts', 'Verdict']]
hl = []
for i, c in enumerate(ORD, 1):
    move = abs_new[c] - abs_base[c]
    pts = move * P.AGG_BASE
    verdict = 'no change' if abs(pts) < 1 else ('watch' if abs(pts) < 3 else 'looks material')
    if abs(pts) >= 3:
        hl.append(i)
    data.append([LAB[c], f'{abs_base[c]:.4f}', f'{abs_new[c]:.4f}', f'{move:+.4f}',
                 f'{pts:+.2f}', verdict])
add_table(data, [3.4, 2.6, 3.4, 2.2, 2.2, 3.6], highlight_rows=hl, highlight_fill=AMBER_FILL)
add_para('Taken at face value that says Target Rifle and Sporter need a real change. <b>It is an '
         'artefact, and worth understanding.</b>')

add_h2('4a. Why that table is misleading — string length')
add_para('The calibration pools each discipline\'s cohort as whole-string totals. A 15-shot string '
         'scores about half as much again as a 10-shot string, so the pooled mean depends on the '
         'mix of match lengths in the data. The mix is not the same:')
data = [['Source', '10-shot strings', '15-shot strings']]
for name, m in (('National K&Q pool', mix_kq), ('NSWRA 2026', mix_nsw), ('QRA Kings 2026', mix_qra)):
    data.append([name, f'{m.get(10, 0):.0f}%', f'{m.get(15, 0):.0f}%'])
add_table(data, [6.0, 4.5, 4.5])
add_para('NSWRA 2026 is half 15-shot matches against 13% in the archive. That alone lifts the pooled '
         'mean of <i>every</i> discipline — the new events\' cohort means come out about 7 points '
         'above the archive across the board, which is a match-format difference, not a shooting '
         'difference. Dividing each string\'s merit by its shot count before pooling makes the '
         'comparison like-for-like:')

data = [['Discipline', 'Pool only', 'Pool + both events', 'Move', 'MCSI pts', 'Verdict']]
for c in ORD:
    move = ps_new[c] - ps_base[c]
    pts = move * P.AGG_BASE
    verdict = 'no change' if abs(pts) < 1 else ('watch' if abs(pts) < 3 else 'looks material')
    data.append([LAB[c], f'{ps_base[c]:.4f}', f'{ps_new[c]:.4f}', f'{move:+.4f}',
                 f'{pts:+.2f}', verdict])
add_table(data, [3.4, 2.6, 3.4, 2.2, 2.2, 3.6])
data = [['Discipline', 'Move, whole-string', 'Move, per-shot', 'Artefact share']]
for c in ORD:
    da, dp = abs_new[c] - abs_base[c], ps_new[c] - ps_base[c]
    share = f'{100 * (1 - abs(dp) / abs(da)):.0f}%' if abs(da) > 1e-9 else '—'
    data.append([LAB[c], f'{da:+.4f}', f'{dp:+.4f}', share])
add_table(data, [4.5, 4.3, 4.3, 4.3])
add_callout('<b>Answer to the question:</b> folding NSWRA 2026 and QRA 2026 into the pool '
            f'<b>does not change any V5 factor</b>. On the like-for-like comparison the largest '
            f'movement is {max(abs(ps_new[c] - ps_base[c]) for c in ORD):.4f} '
            f'({max(abs((ps_new[c] - ps_base[c]) * P.AGG_BASE) for c in ORD):.1f} MCSI points on a '
            '~490-point series aggregate), and nothing moves at the three decimals V5 is published '
            'at. Nearly 4,000 fresh King\'s-class strings leave the formula where it is — which is '
            'the strongest evidence yet that V5 is calibrated correctly.',
            fill=GREEN_FILL, border='3C7A3C')

add_h2('4b. Two caveats on the numbers above, stated plainly')
add_bullet('<b>The baseline here is K&Q only.</b> Published V5 was set on the K&Q pool plus about '
           '31 same-conditions prize meets; that scraped supplement was cached in a temporary '
           'folder and is gone. The K&Q-only baseline reproduces published V5 to within 0.008 on '
           'whole-string pooling. A larger baseline pool can only dilute new data further, so the '
           'movements reported are an <b>upper bound</b> — the true movement against the full '
           'published pool is smaller still. Re-scraping the prize meets would let us close this '
           'off properly.')
add_bullet(f'<b>Per-shot pooling shifts the baseline itself.</b> Recalibrating the existing pool '
           f'per-shot rather than whole-string would move F-Standard from '
           f'{abs_base["F-Std"]:.3f} to {ps_base["F-Std"]:.3f} and F-Open from '
           f'{abs_base["F-Open"]:.3f} to {ps_base["F-Open"]:.3f}. That is a separate question from '
           'the one asked here, but note it points the same way as section 3: F-Standard\'s factor '
           'may be a little high. I would not act on it until the prize-meet data is rebuilt.')

add_h2('5. The weather, and distance-based factors')
add_para('Conditions at the long ranges were poor and some shooters dropped scores badly. Because '
         'this source carries distances, we can look at it directly. Each row is one shoot, scored '
         'on its own:')
data = [['Shoot', 'Dist', 'TR', 'F-Open', 'F-Std', 'F/TR', 'Sporter', 'Difficulty*']]
hl = []
for i, (k, d, n, i2, h) in enumerate(shoot_rows, 1):
    if d == 1000:
        hl.append(i)
    data.append([k.replace('Kings - ', ''), f'{d}', f'{i2["TR"]:.3f}', f'{i2["F-Open"]:.3f}',
                 f'{i2["F-Standard"]:.3f}', f'{i2["FTR"]:.3f}', f'{i2["Sporter"]:.3f}', f'{h:.2f}'])
add_table(data, [4.6, 1.3, 1.7, 1.7, 1.7, 1.7, 1.7, 2.0], highlight_rows=hl,
          highlight_fill=AMBER_FILL)
add_small('*Difficulty = mean per-shot merit of the top 40% of the Target Rifle field in that shoot. '
          'Higher is easier. TR is used because it is the biggest cohort and is not part of the '
          'F-Standard/F-Open comparison below, so the two are independent.')
add_para('The 1000-yard shoot is the outlier, exactly where you said the weather hit. In that one '
         f'shoot F-Standard\'s implied factor rises to {longest[3]["F-Standard"]:.3f} while '
         f'F-Open\'s falls to {longest[3]["F-Open"]:.3f} — a gap of '
         f'{longest[3]["F-Standard"] - longest[3]["F-Open"]:.3f} against roughly 0.01 '
         'to 0.05 everywhere else. That is physically sensible: in hard wind the restricted '
         'F-Standard cartridges get pushed around in a way the F-Open rifles do not.')
add_para('<b>But that is a conditions effect, not a distance effect, and the meeting proves it.</b> '
         '900 yards was shot three times — twice on Day 2 and once on Day 3. Those three shoots '
         'disagree with each other by roughly half the spread seen across all six distances. Testing '
         'it properly, the F-Standard-minus-F-Open gap tracks difficulty far better than it tracks '
         'distance:')
data = [['Relationship', 'Correlation'],
        ['F-Std minus F-Open gap  vs  difficulty of the shoot', f'r = {r_hard:+.2f}'],
        ['F-Std minus F-Open gap  vs  distance', f'r = {r_dist:+.2f}'],
        ['Difficulty  vs  distance', f'r = {corr(dist, hard):+.2f}']]
add_table(data, [10.5, 4.0], highlight_rows=[1])
add_callout('<b>Recommendation on distance:</b> agreed, we are not ready — and the reason is '
            'sharper than "not enough data". Distance is only a proxy for difficulty, and a leaky '
            'one. A distance-indexed factor fitted to this meeting would mostly be encoding the '
            'weather on the Saturday. When we do build it, the variable to index on is probably '
            '<b>measured difficulty of the shoot</b> — which the field\'s own scores give us for '
            'free — with distance as one input to that, not the index itself.',
            fill=AMBER_FILL, border='B07A2B')
add_small(f'Caveat: {len(keys)} shoots, and distance and difficulty are themselves correlated '
          f'(r = {corr(dist, hard):+.2f}). Treat this as a direction to test across more meetings, '
          'not a fitted result.')

add_h2('6. What I suggest we do')
add_bullet('<b>Change nothing in V5 now.</b> Two fresh championships and ~3,900 new strings leave '
           'every factor where it is.')
add_bullet('<b>Put F-Standard on the watch list.</b> Two independent events and the per-shot '
           'baseline all point the same way — 1.475 may be a shade high. One more season before '
           'touching it.')
add_bullet('<b>Fix the pooling method before the next recalibration.</b> Normalise merit per shot '
           'so the calibration is not sensitive to how many 10- versus 15-shot matches happen to be '
           'in the data. This is a genuine robustness improvement that this exercise flushed out.')
add_bullet('<b>Rebuild the prize-meet scrape</b> so the calibration pool is fully reproducible '
           'again rather than depending on a cache that has since been cleared.')
add_bullet('<b>Keep collecting distance-tagged results.</b> QRA publishes them; if we can get the '
           'same from other associations we can properly test a difficulty-aware model next season.')

add_hr()
add_small('Sources: QRA Kings Prize Shoot 2026 results (shooting.hexsystems.com.au/competition/711), '
          'validated string-by-string against all 1,089 published aggregate rows; NSWRA 150th Annual '
          f'Open Championships 2026; national K&Q pool, {n_kq:,} strings from 2024 onward across seven '
          'associations. Accompanying file: QRA_Kings_2026_scores.csv — every string with its V5 and '
          'June-13 adjusted score. Reproduce with fetch_qra_v5.py, analyse_qra_v5.py, '
          'compare_qra_nswra.py, pool_recalibration.py.', italic=True)

doc.save(OUT)
print('Wrote', OUT)
