"""One-page 'why V5 is fairer' argument, framed on OUTCOMES not formula mechanics.
Head-to-head numbers come from GRC_MCSI_Club_Champion_Master_2026_v5_with_peter.xlsx
(Peter = 2014-lineage formula vs V5 on the same current club results):
  Spearman rank corr 0.987 ; top-10 overlap 9/10 ; 36 shooters.
  Avg rank move Peter->V5: FO -2.4, FS -1.0, FTR 0.0, SporterO +1.3, SporterP 0.0, TR +1.0.
Same house style as build_gerald_response_docx.py.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/dancomerford/Desktop/claude/nraadatabase/GRC_MCSI_V5_Why_Fairer_1page.docx'

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
GREEN = RGBColor(0x2E, 0x6B, 0x2E)
HEADER_ROW_FILL = 'EEF3F8'
ROW_HIGHLIGHT = 'FFF9E6'
GREY_FILL = 'F4F4F4'
GREEN_FILL = 'EAF3EA'
GREEN_BORDER = '3C7A3C'

doc = Document()
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _left_border(cell, hex_color, sz='24'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), sz)
    left.set(qn('w:space'), '0'); left.set(qn('w:color'), hex_color)
    borders.append(left); tcPr.append(borders)


def _add_runs(paragraph, text):
    text = text.replace('&amp;', '&')
    pattern = re.compile(r'(<b>.*?</b>|<i>.*?</i>)', re.S)
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith('<b>') and part.endswith('</b>'):
            paragraph.add_run(part[3:-4]).bold = True
        elif part.startswith('<i>') and part.endswith('</i>'):
            paragraph.add_run(part[3:-4]).italic = True
        else:
            paragraph.add_run(part)


def add_para(text, *, size=10, bold=False, italic=False, color=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
        if bold: r.bold = True
        if italic: r.italic = True
        if color is not None: r.font.color.rgb = color
    return p


def add_title(text):
    add_para(text, size=16, bold=True, color=NAVY, space_after=2)


def add_subtitle(text):
    add_para(text, size=10, color=GREY, space_after=5)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(11.5); r.bold = True; r.font.color.rgb = NAVY


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10)


def add_small(text, italic=True):
    add_para(text, size=8.5, italic=italic, color=LIGHT_GREY, space_after=5)


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); p._p.get_or_add_pPr().append(pBdr)


def add_callout(text, fill=HEADER_ROW_FILL, border=NAVY):
    border_hex = border if isinstance(border, str) else '1F4E78'
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0); cell.width = Cm(17.4)
    _shade_cell(cell, fill); _left_border(cell, border_hex)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    _add_runs(p, text)
    for r in p.runs:
        r.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_table(data, col_widths_cm=None, highlight_rows=None):
    n_rows, n_cols = len(data), len(data[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False; table.style = 'Table Grid'
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for r_idx, row_data in enumerate(data):
        is_header = r_idx == 0
        is_highlight = highlight_rows and r_idx in highlight_rows
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
            if col_widths_cm and c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(p, str(text))
            for r in p.runs:
                r.font.size = Pt(9.5)
                if is_header:
                    r.bold = True; r.font.color.rgb = NAVY
            if is_header:
                _shade_cell(cell, HEADER_ROW_FILL)
            elif is_highlight:
                _shade_cell(cell, ROW_HIGHLIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ── Header ───────────────────────────────────────────────────────────────────
add_title('Why V5 is fairer — and why it barely changes who wins')
add_subtitle('Geelong Rifle Club · MCSI · the argument in terms of the final result, not the '
             'formula · July 2026')
add_hr()

# ── 1 ─────────────────────────────────────────────────────────────────────────
add_h2('What we have been scoring on is a 2014 system')
add_para(
    'The club’s current formula (Peter’s) is a faithful re-build of the original Mixed Category '
    'Score Index published in 2014 (ozfclass.com). It uses the same structure and even ranks the '
    'disciplines’ difficulty in the same order. In effect, <b>we have been scoring the championship '
    'against a snapshot of the sport as it was a decade ago.</b>')

# ── 2 ─────────────────────────────────────────────────────────────────────────
add_h2('Why that is now unfair')
add_para(
    'Every shooter is being measured against how people shot in <b>2014</b>. Two things have '
    'changed since: <b>equipment has moved on</b> — most visibly in F-Open, where scores today are '
    'far higher than in 2014 — and <b>Sporter was never part of the 2014 build</b>, so it has been '
    'carried on an estimate rather than real data. A shooter today is being judged against the '
    'wrong yardstick.')

# ── 3 ─────────────────────────────────────────────────────────────────────────
add_h2('What V5 changes')
add_para(
    'V5 is the <b>same index, recalibrated on ~19,800 recent (2024-onward) King’s/Queen’s scores</b> '
    '— so each shooter is measured against <b>today’s</b> real competition, and Sporter is '
    'calibrated from actual results for the first time. Nothing about the intent changes; only the '
    'yardstick is brought up to date.')

# ── 4 — the outcome ────────────────────────────────────────────────────────────
add_h2('The final answer: V5 agrees with the system we trust — except where 2014 is stale')
add_para('Running both formulas over the <b>same current club results</b> (36 shooters):')
add_callout(
    '<b>The two systems rank the club almost identically — a rank correlation of 0.987, and 9 of '
    'the top 10 are the same people.</b> V5 is not a re-invention; it keeps everything the current '
    'system already gets right.', fill=GREEN_FILL, border=GREEN_BORDER)
add_para('The only material movements fall in the two disciplines the 2014 data cannot represent:')
add_table([
    ['Discipline', 'Placing under V5', 'Why it moves'],
    ['F-Open', '↓ 2–3 places', 'Modern equipment scores far higher than in 2014; recent data '
     'stops over-crediting it.'],
    ['Sporter (Open)', '↑ ~1 place', 'Calibrated from real data for the first time — no longer an '
     'estimate.'],
    ['Target Rifle', '↑ ~1 place', 'Small correction against current scoring.'],
    ['F-Standard', '↓ ~1 place', 'Small correction against current scoring.'],
    ['F/TR · Sporter (Prod.)', 'unchanged', 'Already well represented — essentially no change.'],
], col_widths_cm=[4.4, 3.2, 9.8], highlight_rows={1, 2})
add_small('Source: same club results scored under Peter’s (2014-lineage) factors and under V5. '
          'A correction of a place or two can occur at the very top — which is exactly where '
          'getting the yardstick right matters most.')

# ── 5 — conclusion ─────────────────────────────────────────────────────────────
add_hr()
add_callout(
    '<b>In one line:</b> V5 gives the same result as the system we already trust, everywhere '
    'except where that system depends on decade-old 2014 data — and there it gives a fairer one, '
    'based on how members actually shoot today.')

doc.save(OUT)
print(f'Wrote {OUT}')
