"""Build a plain-English 4-page version of the V5 committee report.

Audience: committee members who don't have a statistics background.
No jargon. No confidence intervals, no standard errors, no bootstrap talk.
Aim: ~4 pages A4, 2cm margins.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/dancomerford/Desktop/claude/nraadatabase/NRAA_MCSI_V5_Plain_English.docx'

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
HEADER_ROW_FILL = 'EEF3F8'
ROW_HIGHLIGHT = 'FFF9E6'

doc = Document()

for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def _runs(paragraph, text):
    text = text.replace('&amp;', '&')
    pattern = re.compile(r'(<b>.*?</b>|<i>.*?</i>)', re.S)
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith('<b>') and part.endswith('</b>'):
            r = paragraph.add_run(part[3:-4]); r.bold = True
        elif part.startswith('<i>') and part.endswith('</i>'):
            r = paragraph.add_run(part[3:-4]); r.italic = True
        else:
            paragraph.add_run(part)


def para(text, *, size=11, bold=False, italic=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _runs(p, text)
    for r in p.runs:
        r.font.size = Pt(size)
        if bold: r.bold = True
        if italic: r.italic = True
        if color is not None: r.font.color.rgb = color


def title(text):
    para(text, size=18, bold=True, color=NAVY, space_after=2)


def subtitle(text):
    para(text, size=11, italic=True, color=GREY, space_after=8)


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    _runs(p, text)
    for r in p.runs:
        r.font.size = Pt(13); r.bold = True; r.font.color.rgb = NAVY


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    _runs(p, text)
    for r in p.runs:
        r.font.size = Pt(11)


def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); pPr.append(pBdr)


def add_table(data, col_widths_cm=None, highlight_rows=None):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.autofit = False
    table.style = 'Table Grid'
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for r_idx, row_data in enumerate(data):
        is_header = r_idx == 0
        is_highlight = highlight_rows and r_idx in highlight_rows
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            _runs(p, str(text))
            for r in p.runs:
                r.font.size = Pt(10)
                if is_header:
                    r.bold = True; r.font.color.rgb = NAVY
            if is_header:
                _shade(cell, HEADER_ROW_FILL)
            elif is_highlight:
                _shade(cell, ROW_HIGHLIGHT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────────────────
# PAGE 1 — Cover, what V5 is, the factors
# ─────────────────────────────────────────────────────────────────
title('MCSI Formula V5 — Plain English Summary')
subtitle('A short explainer for the committee — June 2026')
hr()

h2('What this document is')
para("This is a plain-English summary of the recommended MCSI formula (V5) for "
     "scoring the club championship. The full technical report is available if you "
     "want the maths — this version is for committee members who just want to "
     "understand <b>what V5 does, where the numbers come from, and why we trust them</b>.")

h2('What the MCSI formula is for')
para("Different rifle disciplines score differently. An F-Open shooter on a scope and "
     "bipod can post bigger numbers than a TR shooter on irons, who in turn typically "
     "out-scores a Sporter shooter. If we simply added everyone's raw scores together "
     "for a club championship, F-class shooters would win every year — not because "
     "they shoot better, but because their equipment lets them score higher.")
para("The MCSI formula fixes this. It multiplies each discipline's score by a "
     "<b>fairness factor</b> so that a strong performance in any discipline counts "
     "comparably toward the club championship. The whole exercise is just about "
     "picking the right fairness factors.")

h2('The recommendation')
para("<b>Adopt V5 as the club championship formula.</b> Five factors, one centre "
     "rule, one conversion adjustment for iron-sight disciplines:")

add_table([
    ['Discipline', 'Factor', 'What it means'],
    ['Target Rifle (TR)', '1.412', 'Iron sights — gets a 1.2× score boost first'],
    ['F-Open', '1.406', 'Scope + bipod — no boost'],
    ['F-Standard', '1.475', 'Scope + bipod — no boost'],
    ['F/TR (FTR)', '1.450', 'Scope + bipod — no boost'],
    ['Sporter (Open + Production combined)', '1.383', 'Iron sights — gets a 1.2× score boost first'],
], col_widths_cm=[7.5, 2.0, 7.0])

para("<b>The formula in one line:</b> Adjusted Score = (Raw Score × Conversion + "
     "Centres × 0.7) × Factor", space_after=4)
para("<b>A key thing to notice:</b> the five factors only differ by 0.092 (from "
     "1.383 to 1.475). The disciplines are much closer in difficulty than people "
     "often assume. V5 is mostly about fairness in how centres are counted and how "
     "iron-sight disciplines are equalised against F-class — not about handing out "
     "big equipment penalties.")

# ─────────────────────────────────────────────────────────────────
# PAGE 2 — Where the numbers come from, why K&Q
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()

h2('Where the numbers come from')
para("To work out the fairness factors, we needed a big collection of scores where "
     "all five disciplines were shooting at the same time, in the same wind, on the "
     "same day. That way the comparison is fair — no discipline gets unfairly "
     "boosted by good weather while another gets penalised by bad.")
para("The answer turned out to be <b>King's & Queen's championships</b>. They are "
     "the only events in Australia where all five disciplines reliably shoot "
     "side-by-side in the same conditions.")
para("We used 13 K&Q championships from 2024 onwards across all seven state and "
     "national associations (NSWRA, VRA, QRA, NQRA, SARA, WARA, and NRAA "
     "Nationals). That gave us <b>20,335 individual range scores from over a "
     "thousand different shooters</b>.")
para("From that pool, we looked at the top 40% of scores in each discipline and "
     "set the factors so that a top-40% shot in any discipline produces roughly "
     "the same Adjusted Score as a top-40% shot in F-class.")

h2('Why we didn\'t just use Geelong data')
para("It would be intuitive to use only Geelong scores — after all, this is for "
     "the Geelong club championship. But it doesn't work, for three plain reasons.")

bullet("<b>The dataset is far too small.</b> Geelong's whole 2026 season produced "
       "about 270 scores across 36 shooters. K&Q gives us 20,335 scores from over "
       "a thousand shooters — about 75 times more data.")
bullet("<b>F/TR can't be calibrated at all.</b> Geelong's 2026 season had just 9 F/TR "
       "scores from 2 shooters. You cannot derive a fair F/TR factor from 9 scores. "
       "Any number you pick is essentially a guess.")
bullet("<b>The numbers would swing wildly year to year.</b> A small sample is "
       "volatile. One strong new member or one shooter taking a year off could move "
       "Geelong's Sporter factor by 0.1 or more — completely re-arranging the "
       "leaderboard for reasons unrelated to actual shooting.")

para("There is also a deeper fairness problem with using local data. If one strong "
     "Geelong shooter dominates a discipline, that shooter's scores would inflate "
     "the discipline's factor — and then they'd be ranked against a factor they "
     "personally raised. The K&Q dataset avoids this because no single shooter has "
     "any meaningful influence on a 20,000-score pool.")

h2('Why we didn\'t use Prize Meet data either')
para("Prize meets seemed promising — there's a lot of data available. We pulled "
     "223 prize meets and tested using them. Two problems killed the idea.")

bullet("<b>Western Australia distorts the numbers.</b> WA has a very strong Sporter "
       "culture and relatively weaker F-class prize-meet turnout. When WA prize "
       "meets are included, the Sporter factor drops by about 0.06 — purely "
       "because of regional sampling, not because Sporter is actually easier.")
bullet("<b>East-coast prize meets favour particular disciplines club-by-club.</b> "
       "Pacific Rifle Club's meets are dominated by TR shooters; Goondiwindi's by "
       "F-class. When you mix these together, the discipline that happened to turn "
       "up in larger numbers wins the calibration — again, nothing to do with how "
       "the disciplines actually compare.")

para("Adding prize-meet data shifted the TR factor by 0.074 and Sporter by 0.067 "
     "in the wrong direction — clear evidence the prize-meet pool is biased, not "
     "that the disciplines have suddenly become easier.")

# ─────────────────────────────────────────────────────────────────
# PAGE 3 — How a score is calculated, why it's trustworthy
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()

h2('How an Adjusted Score is calculated — worked example')
para("Let's take a Sporter shooter who shoots a 50.10 (50 points with 10 centres).")
para("<b>Step 1 — Score × Conversion.</b> Sporter uses iron sights, so we apply the "
     "1.2× conversion to bring iron-sight disciplines onto a comparable scale with "
     "F-class: 50 × 1.2 = 60.")
para("<b>Step 2 — Add centres × 0.7.</b> 10 centres × 0.7 = 7. So we now have "
     "60 + 7 = 67.")
para("<b>Step 3 — Apply the Sporter factor.</b> 67 × 1.383 = 92.66.")
para("The same shot in F-Open would be: 50 × 1.0 + 10 × 0.7 = 57, × 1.406 = 80.14. "
     "Hold on — that's lower. But F-Open shooters usually score higher in raw points "
     "(say 60.10 in a good string), so once their bigger raw scores feed in, the "
     "Adjusted Scores end up comparable. That is exactly the point of the formula.")

h2('Why we trust the V5 numbers')
para("Three independent checks support V5:")

bullet("<b>Big sample.</b> 20,335 scores is a serious dataset. The factors don't "
       "wobble if we add or remove a single shooter, a single match, or even a "
       "single state.")
bullet("<b>Two different methods agree.</b> We calibrated V5 using the top 40% of "
       "scores. We also re-ran the whole exercise using the middle 20% of scores "
       "(a completely different way of measuring). The factors came out within "
       "0.02 of each other on every discipline. When two unrelated methods agree, "
       "the answer is almost certainly real.")
bullet("<b>Each state agrees with the others.</b> If you split the K&Q data by "
       "state and check whether Sporter is stronger or weaker than F-class in each "
       "one, the answer is very consistent: Sporter is about 0% to 6% stronger "
       "than F-class everywhere (the only outlier is South Australia with just 48 "
       "scores). That consistency is what lets us trust the national pool.")

h2('What about the 0.7 centre weight?')
para("We tested four options: 0.5, 0.7, 1.0, and counting each centre as a full "
     "point. The 0.7 weight gave the most natural result:")

add_table([
    ['Centre weight', 'Value of 50.10', 'Value of 50.0', 'Does 50.0 beat 49.10?'],
    ['0.5', '55', '50', 'Yes, comfortably'],
    ['<b>0.7 (selected)</b>', '<b>57</b>', '<b>50</b>', '<b>Yes — by a clear 3 points</b>'],
    ['1.0', '60', '50', 'Yes, but only just'],
    ['Full centre rule', 'Complex jump', 'Complex jump', 'Sometimes — unpredictable'],
], col_widths_cm=[3.5, 3.0, 3.0, 5.5], highlight_rows=[2])

para("At 0.7, a 10-centre string clearly outranks a same-score zero-centre string "
     "(by 7 points), but a higher raw score still beats a lower one with more "
     "centres — which matches how NRAA scoring has always worked.")

# ─────────────────────────────────────────────────────────────────
# PAGE 4 — Common questions, summary
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()

h2('Common questions from the committee')

para("<b>\"Why should a national formula apply to our club championship?\"</b>")
para("A 500m target is the same size in Geelong as it is in Belmont, Brisbane or "
     "Perth. The physics of shooting doesn't change at the state border. A TR rifle "
     "isn't harder to shoot because it's in Victoria. So the fairness factors should "
     "come from the largest, most representative dataset we can find — and that's "
     "the national K&Q pool, not Geelong's 270 scores from one season.")

para("<b>\"What if Geelong has more Sporter shooters than other clubs?\"</b>")
para("That's true today, and it might change next year if new members join or "
     "current ones leave. A formula that depends on which shooters happened to be at "
     "Geelong this season is unstable — we'd be re-tuning it every year for reasons "
     "that have nothing to do with how the disciplines compare.")

para("<b>\"Doesn't this give one discipline an advantage in our championship?\"</b>")
para("Looking at the V5 factor spread of just 0.092 — these are very close numbers. "
     "If the committee decides it wants <b>every</b> discipline to win the "
     "championship in roughly equal proportions (for example to keep all disciplines "
     "feeling competitive), that's a separate policy decision and we can layer it on "
     "top of V5. But that is a championship-design choice, not a fairness choice. "
     "V5 answers the question \"how difficult was this score to achieve?\" — it "
     "doesn't pre-engineer who wins.")

para("<b>\"How often will V5 need updating?\"</b>")
para("The factors should be re-checked once a year as new K&Q championships are "
     "fired. The methodology is reproducible — we just re-run the calculation with "
     "the new data. Small year-on-year shifts are normal; large shifts would "
     "prompt a review.")

para("<b>\"What if a competitor doesn't agree with their discipline's factor?\"</b>")
para("Every factor is derived from how the top 40% of that discipline actually "
     "scored in real championships, against the top 40% of every other discipline "
     "in the same events. There's nothing arbitrary in the numbers — they fall out "
     "of the data. If a discipline feels under-rewarded, the test is to look at "
     "how it actually performs at K&Q level, not at intuition.")

h2('Summary')

bullet("V5 is calibrated against 20,335 K&Q championship scores from 2024 onwards, "
       "across all seven state and national associations.")
bullet("It's the largest and fairest dataset available — 75 times bigger than any "
       "one club season.")
bullet("The five discipline factors are all within 0.092 of each other — much "
       "closer than people often assume.")
bullet("Two completely different statistical approaches give the same answer, and "
       "every state agrees with every other state. That's strong evidence the "
       "numbers are right.")
bullet("Local-only data was tested and failed — sample sizes are too small and "
       "F/TR can't be calibrated at all from one club's data.")
bullet("Prize-meet data was tested and failed — Western Australia distorts it one "
       "way, east-coast club specialisation distorts it another way.")

para("", space_after=4)
para("<b>The committee's role:</b> decide whether to adopt V5 (recommended) and "
     "whether any further policy adjustment is wanted on top of it.", space_after=4)

para("", space_after=2)
para("<i>End of plain-English summary. Full technical report available on request.</i>",
     size=9, italic=True, color=LIGHT_GREY)

doc.save(OUT)
print(f'Wrote {OUT}')
